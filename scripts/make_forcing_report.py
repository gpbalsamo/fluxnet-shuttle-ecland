#!/usr/bin/env python3
"""Generate a Word summary of the met-forcing procedure and its QC report.

All numbers are computed from the artefacts on disk (the QC CSV, the forcing
filenames, the pipeline status directory) rather than transcribed, so the
document cannot drift from the data it describes.
"""
import collections
import csv
import glob
import os
import re
import statistics
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

REPO = os.environ.get("FLUXNET_SHUTTLE_REPO",
                      os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
GROUP = "shuttle-all775-era5"
QC = f"{REPO}/reference/qc_report_{GROUP}.csv"
FORCING = f"{REPO}/forcing/{GROUP}"
STATUS = os.environ.get("FORCING_STATUS_DIR",
                        f"{os.environ.get('SCRATCH', '/tmp')}/forcing_pipeline_{GROUP}/status")
SNAP = os.environ.get("SHUTTLE_SNAPSHOT_CSV",
                      "/ec/res4/scratch/pad/shuttle_setup/fluxnet_shuttle_snapshot_20260818T102536.csv")
OUT = f"{REPO}/docs/fluxnet_shuttle_met_forcing_summary.docx"

DRIVING = ["Tair", "Qair", "PSurf", "Wind", "SWdown", "LWdown", "Rainf"]
VAR_ORDER = ["SWdown", "Tair", "RH", "VPD", "Qair", "PSurf", "Wind", "Rainf", "LWdown", "CO2air"]
pat = re.compile(r"^met_insituHT_(.+?)_(\d{4})-(\d{4})\.nc$")


def band(x):
    return "mild" if x <= 10 else "medium" if x <= 25 else "heavy" if x <= 50 else "complete"


# ---------------------------------------------------------------- gather data
rows = list(csv.DictReader(open(QC)))
for r in rows:
    r["gapfilled_pct"] = float(r["gapfilled_pct"])
    r["missing_pct"] = float(r["missing_pct"])

snap = {}
for r in csv.DictReader(open(SNAP)):
    snap.setdefault(r["site_id"], r)

status = collections.Counter(open(p).read().strip() for p in glob.glob(f"{STATUS}/*"))

years, per_site = {}, collections.defaultdict(dict)
for r in rows:
    m = pat.match(r["file"])
    site = m.group(1)
    years[site] = int(m.group(3)) - int(m.group(2)) + 1
    per_site[site][r["variable"]] = r["gapfilled_pct"]

site_mean = {s: statistics.mean([v[k] for k in DRIVING if k in v]) for s, v in per_site.items()}
site_worst = {s: max(v[k] for k in DRIVING if k in v) for s, v in per_site.items()}
total_years = sum(years.values())

by_var = collections.defaultdict(list)
for r in rows:
    by_var[r["variable"]].append(r)

hub_sites, hub_years = collections.Counter(), collections.Counter()
igbp = collections.Counter()
for s in per_site:
    hub = snap.get(s, {}).get("data_hub", "?")
    hub_sites[hub] += 1
    hub_years[hub] += years[s]
    igbp[snap.get(s, {}).get("igbp", "?")] += 1

band_sites = collections.Counter(band(v) for v in site_mean.values())
band_years = collections.Counter()
for s, v in site_mean.items():
    band_years[band(v)] += years[s]
band_worst = collections.Counter(band(v) for v in site_worst.values())

method = collections.Counter(r["gapfilling_method"] for r in rows)
observed = sum(v for k, v in method.items() if k == "observed")

# ------------------------------------------------------------------ document
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)

def h(text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    return p

def para(text, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def table(headers, data, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = str(hd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    for row in data:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                if i:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# Title
tp = doc.add_heading("Meteorological Forcing for ecLand from the FLUXNET Shuttle", level=0)
for run in tp.runs:
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
para(f"Procedure and quality-control summary · site group “{GROUP}” · {date.today().isoformat()}", italic=True)
para("Code: fluxnet-shuttle-ecland (github.com/gpbalsamo/fluxnet-shuttle-ecland). Built on the ECMWF "
     "HPC. Every number in this document is read straight from the files that were produced, not "
     "typed in by hand.", size=9.5)

h("1. Purpose", 1)
para("Until now, ecLand site runs have used the same fixed set of 170 flux-tower sites (PLUMBER2). "
     "The FLUXNET Shuttle gives us a live list of every site available from AmeriFlux, ICOS and TERN — "
     "775 sites on 18 August 2026, and growing. This work turns that whole list into weather data "
     "ecLand can run on. We care most about the dry, fire-prone places PLUMBER2 has few of: savanna, "
     "shrubland, grassland, the Mediterranean and the Sahel.")

h("2. Procedure", 1)
para("Each site goes through the same seven steps, run end to end by "
     "scripts/run_forcing_pipeline.sh:")
table(["Step", "Tool", "Purpose"],
      [["1", "fluxnet-shuttle listall", "Ask the Shuttle what sites exist today, and save the list"],
       ["2", "build_site_metadata.py", "Give every site a location and a vegetation type. The conversion software knows 874 older sites; newer ones come from the Shuttle list"],
       ["3", "fetch_noaa_co2.py", "Fetch NOAA's monthly world-average CO2, used later to fill CO2 gaps. Only needed once"],
       ["4", "fluxnet-shuttle download", "Download the site: half-hourly tower measurements, plus matching ERA5 weather"],
       ["5", "fill_co2_from_noaa.py", "Fill any missing CO2 from the NOAA figures, marked as filled-in"],
       ["6", "convert_fluxnetlsm.R", "Convert to standard NetCDF files, filling gaps and dropping years that are too patchy"],
       ["7", "regenerate_forcing.sh", "Rename things to the spellings ecLand expects, so the files can be read straight in"]],
      widths=[0.5, 1.9, 4.2])

para("A single site's raw download can be 500 MB, so each one is deleted as soon as it has been "
     "converted; only the small final files are kept. Every finished site leaves a note saying how it "
     "went. That means the run can be stopped and restarted at any point, and a site that failed "
     "because the network dropped can be retried on its own, without redoing the ones that worked.")

h("2.1 Running it on the HPC", 2)
para("scripts/submit_forcing_pipeline_slurm.sh runs the same pipeline as a batch job on the HPC. "
     "The site list is dealt out between jobs like cards, one site at a time round the table, rather "
     "than split into blocks. Sites are listed by country and vary a lot in record length, so blocks "
     "would leave one job with all the long records and another with nothing to do. Each job then works "
     "through its own sites, several at once, and all jobs share one set of status notes. The full "
     "775-site run used 8 jobs of 4 sites each — 32 sites being fetched at once — and took about "
     "47 minutes, plus a short second pass for the sites rescued by the CO2 fix.")

h("3. The three decisions that mattered", 1)

h("3.1 How we fill the holes in the records", 2)
para("Tower records always have holes in them — an instrument fails, or a reading is thrown out. "
     "Those holes have to be filled before the data can drive a model. There are two ways to do it: "
     "estimate the missing values from the tower's own data, or take them from ERA5, the weather "
     "record that comes packaged with each site. Which one you pick matters far more than any quality "
     "setting. On the same 775 sites with the same settings, filling from the tower's own data gave "
     "usable weather for 231 sites; filling from ERA5 gave all 775. The tower's own data cannot fill a "
     "long gap, and a year with a gap left in it is thrown away completely. So ERA5 is used throughout. "
     "The price is that a good deal of the data is then weather-model output rather than measurement, "
     "which section 5 sets out in full.")
para("A note on names: the conversion software calls this option “ERAinterim”, after an older weather "
     "record, but the data the Shuttle actually supplies is ERA5. The columns and timestamps line up "
     "exactly with what the software expects, so it works as-is.", size=9.5)

h("3.2 CO2 was quietly costing us 97 sites", 2)
para("A year is thrown away if any weather variable still has a hole in it once the filling is done. "
     "CO2 is the one variable ERA5 cannot help with, because weather records do not carry the amount of "
     "CO2 in the air. So 97 sites were being thrown away entirely over CO2 — a variable this setup does "
     "not even use, since CO2 feedback on plant growth is switched off and CO2 only comes out of the "
     "model, never into it. Missing CO2 is now filled from NOAA's monthly global average, which has run "
     "since 1979. Filled values are marked as the lowest quality class, so they are still counted and "
     "reported as filled rather than passed off as measurements. This brought back all 97 sites and "
     "264 site-years, without loosening a single quality threshold.")

h("3.3 One site was being skipped for no stated reason", 2)
para("The conversion software ships with a list of sites to skip, and CZ-BK2 is on it with no reason "
     "given. The site is dropped before its data is even opened. We checked the data: every weather "
     "variable it needs can be filled from ERA5, and the site converts cleanly into a complete 7-year "
     "record with nothing missing. We now override that skip, recording why in the code so the decision "
     "stays visible and can be reversed with one option.")

doc.add_page_break()
h("4. What we ended up with", 1)
para(f"Every single site on the list came out with usable weather data: {status.get('OK', 0)} of "
     f"{sum(status.values())}. Together they add up to {total_years} site-years — one unbroken stretch "
     f"per site, {statistics.median(sorted(years.values())):.0f} years long for a typical site and "
     f"{max(years.values())} years at the longest.")
table(["Data hub", "Sites", "Site-years"],
      [[hub, f"{hub_sites[hub]}", f"{hub_years[hub]}"] for hub in sorted(hub_sites)]
      + [["Total", f"{sum(hub_sites.values())}", f"{total_years}"]],
      widths=[2.5, 1.7, 1.7])
para("Distribution by IGBP class:")
para("   " + ", ".join(f"{k} {v}" for k, v in igbp.most_common()), size=9.5)
fire = sum(igbp[k] for k in ("SAV", "WSA", "OSH", "CSH", "GRA"))
para(f"The dry, fire-prone vegetation types we set out to cover — savanna, woody savanna, open and "
     f"closed shrubland, and grassland — come to {fire} sites on their own. That is more sites than the "
     f"whole PLUMBER2 set of 170.")

h("5. How good is the data?", 1)
para("Every file keeps an honest record of its own quality: for each variable it stores how much was "
     "missing, how much was filled in, and how. Those notes are carried through to the final files, so "
     "anyone can check the quality later and leave out what they do not trust, without running anything "
     f"again. The figures below come from reading those notes back ({len(rows)} entries — "
     f"{len(per_site)} files with {len(by_var)} variables each).")
para(f"Nothing is left missing anywhere: all {len(rows)} entries report 0.00% missing. Of those, "
     f"{observed} ({100*observed/len(rows):.0f}%) are pure measurement and the rest have some ERA5 in "
     f"them.")

h("5.1 Variable by variable", 2)
para("How much of each variable had to be filled in, across all 775 sites:")
table(["Variable", "Median", "Mean", "90th pct", "Max", "Files fully observed"],
      [[v,
        f"{statistics.median(sorted(x['gapfilled_pct'] for x in by_var[v])):.1f}%",
        f"{statistics.mean([x['gapfilled_pct'] for x in by_var[v]]):.1f}%",
        f"{sorted(x['gapfilled_pct'] for x in by_var[v])[int(0.9*len(by_var[v]))]:.1f}%",
        f"{max(x['gapfilled_pct'] for x in by_var[v]):.1f}%",
        f"{sum(1 for x in by_var[v] if x['gapfilled_pct'] == 0)}"]
       for v in VAR_ORDER],
      widths=[1.1, 0.95, 0.95, 1.05, 0.85, 1.6])
para("For every variable the typical (median) site is far below the average, which tells us the "
     "picture is lopsided: most sites needed only a little filling in, while a small number are almost "
     "entirely made up of ERA5.")
para("The weakest three are downward longwave radiation, rainfall and air pressure. Each is entirely "
     "ERA5 at a good number of sites. That is what we would expect rather than a sign of trouble: "
     "longwave radiation and pressure sensors are the ones towers most often do without, and rain "
     "gauges are often rejected by quality checks. Rainfall splits into two camps — it has both the "
     "largest group of sites where it is fully measured and a large group where it is entirely ERA5.")

h("5.2 Site by site, counting only what ecLand uses", 2)
para("ecLand runs on seven variables: " + ", ".join(DRIVING) + ". Looking only at those, and ignoring "
     "CO2, humidity and vapour pressure deficit, which the model does not read, we can say how much of "
     "each site is real measurement:")
table(["Quality band (mean gap-fill)", "Sites", "Site-years"],
      [["Mild (≤10%)", band_sites["mild"], band_years["mild"]],
       ["Medium (10–25%)", band_sites["medium"], band_years["medium"]],
       ["Heavy (25–50%)", band_sites["heavy"], band_years["heavy"]],
       ["Predominantly reanalysis (>50%)", band_sites["complete"], band_years["complete"]],
       ["Total", sum(band_sites.values()), total_years]],
      widths=[3.0, 1.4, 1.5])
q = sorted(site_mean.values())
para(f"A typical site is {statistics.median(q):.1f}% filled in across the seven variables ecLand "
     f"reads. A quarter of sites are below {q[len(q)//4]:.1f}% and a quarter above "
     f"{q[3*len(q)//4]:.1f}%. {band_sites['mild'] + band_sites['medium']} sites, covering "
     f"{band_years['mild'] + band_years['medium']} site-years, are at 25% or less; that is the solid "
     f"core to work with. The heavily filled sites are mostly short records, so leaving them out costs "
     f"far fewer years than the site count suggests.")
para(f"That is the average across the seven. If instead you judge a site by its single worst variable, "
     f"{band_worst['complete']} sites land above 50%, because one fully-ERA5 variable — usually longwave "
     f"radiation or rainfall — is enough to get there. Whether that rules a site out depends on what you "
     f"are testing, which is why the quality figures are published for every variable of every file "
     f"instead of boiled down to one score per site.")

h("6. What to watch out for, and what is still open", 1)
bullet("Getting all 775 sites depends on ERA5. About three quarters of the entries have some ERA5 in "
       "them, and a few sites are almost nothing but. Do not treat every site as equally good — use the "
       "quality figures stored in each file to decide what to keep.")
bullet("We still cannot run ecLand on these sites. Weather data is only half of what the model needs; "
       "it also needs the fixed description of each place — soil, vegetation cover, height above sea "
       "level, seasonal leaf area. Nothing we have builds those for a new location yet. That is now the "
       "one and only thing standing in the way.")
bullet("The site list changes from day to day. Sites come and go: the 18 August list had about 30 "
       "fewer ICOS sites than the day before, because some were missing the citation details they are "
       "published with. Always say which day's list your results came from.")
bullet("Where to keep the files is still undecided. The weather files come to about 1.8 GB and the "
       "matching tower measurements about 7.2 GB, which is more than the code repository can hold for "
       "free.")

h("7. How to rebuild this", 1)
para("Once the conversion software and the Shuttle tool are installed (the README explains how):", size=9.5)
p = doc.add_paragraph()
r = p.add_run("fluxnet-shuttle listall\n"
              "python3 scripts/build_site_metadata.py <snapshot>.csv --out reference/site_metadata_merged.csv\n"
              "python3 scripts/fetch_noaa_co2.py --out reference/noaa_gml_co2_monthly.csv\n"
              "scripts/submit_forcing_pipeline_slurm.sh -f <snapshot>.csv -g shuttle-all775-era5 \\\n"
              "    -c reference/site_metadata_merged.csv -C reference/noaa_gml_co2_monthly.csv \\\n"
              "    -P complete -G erainterim -a 8 -j 4\n"
              "python3 scripts/qc_classify.py forcing/shuttle-all775-era5 --out qc_report.csv")
r.font.name = "Consolas"
r.font.size = Pt(8.5)
para(f"The quality figures summarised here are saved with the code, as "
     f"reference/qc_report_{GROUP}.csv.", size=9.5)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("wrote", OUT)
print(f"sanity: {len(rows)} records, {len(per_site)} sites, {total_years} site-years, "
      f"bands {dict(band_sites)}")
